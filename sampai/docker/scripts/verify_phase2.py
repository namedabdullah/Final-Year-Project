"""Phase 2 ingestion acceptance test (run against a live server on :9621).

Hard gate (legacy parser, fast):
  upload .txt -> PROCESSING -> COMPLETED, AI summary generated, Neo4j nodes carry
  the classroom_{id} workspace label, delete clears KB (node count drops) + 404.

External-parser gate (live Docling):
  generate a small .docx -> upload -> COMPLETED via the docling-serve container.
"""

from __future__ import annotations

import io
import sys
import time
import uuid

import httpx
from docx import Document
from neo4j import GraphDatabase

from lightrag.api.sampai.config import SampaiSettings

BASE = "http://127.0.0.1:9621/api/sampai"
_passed = _failed = 0


def check(name, cond, extra=""):
    global _passed, _failed
    if cond:
        _passed += 1; print(f"  PASS  {name}")
    else:
        _failed += 1; print(f"  FAIL  {name}  {extra}")


def neo4j_label_count(cid: int) -> int:
    s = SampaiSettings.load()
    import os
    uri = os.getenv("NEO4J_URI", "bolt://localhost:7687")
    user = os.getenv("NEO4J_USERNAME", "neo4j")
    pw = os.getenv("NEO4J_PASSWORD", "please-change")
    drv = GraphDatabase.driver(uri, auth=(user, pw))
    try:
        with drv.session() as sess:
            rec = sess.run(f"MATCH (n:`classroom_{cid}`) RETURN count(n) AS c").single()
            return rec["c"] if rec else 0
    finally:
        drv.close()


def poll_status(c, headers, fid, timeout=300):
    start = time.time()
    last = None
    while time.time() - start < timeout:
        r = c.get(f"/files/{fid}/status", headers=headers)
        d = r.json()
        cur = f"{d['status']}/{d.get('stage')}"
        if cur != last:
            print(f"      … {cur}")
            last = cur
        if d["status"] in ("completed", "failed"):
            return d["status"]
        time.sleep(2)
    return "timeout"


def main() -> int:
    sfx = uuid.uuid4().hex[:8]
    teacher = {"username": f"t2_{sfx}", "email": f"t2_{sfx}@x.com", "password": "Passw0rd123"}

    with httpx.Client(base_url=BASE, timeout=60) as c:
        c.post("/auth/signup", json=teacher)
        tok = c.post("/auth/login", json={"email": teacher["email"], "password": teacher["password"]}).json()["access_token"]
        H = {"Authorization": f"Bearer {tok}"}

        cid = c.post("/classrooms", headers=H, json={"name": f"Ingest {sfx}"}).json()["id"]
        fid_folder = c.post(f"/folders/classroom/{cid}", headers=H, json={"name": "Materials"}).json()["id"]
        print(f"  classroom={cid} folder={fid_folder}")

        # ── 1) TXT (legacy) hard gate ──
        txt = (
            "Operating Systems Notes. A process is an instance of a running program. "
            "The CPU scheduler selects which process runs next. Deadlock occurs when "
            "processes wait on each other indefinitely. Virtual memory uses paging to "
            "give each process its own address space."
        ).encode()
        r = c.post(f"/files/upload/{fid_folder}", headers=H, files={"upload": ("os_notes.txt", txt, "text/plain")})
        check("txt upload 201", r.status_code == 201, str(r.status_code))
        txt_fid = r.json()["id"]
        print("  ingesting txt (legacy)…")
        st = poll_status(c, H, txt_fid, timeout=180)
        check("txt -> completed", st == "completed", st)
        detail = c.get(f"/files/{txt_fid}", headers=H).json()
        check("txt summary generated", bool(detail.get("description")), str(detail.get("description"))[:60])

        nodes_after = neo4j_label_count(cid)
        check("Neo4j has classroom workspace nodes", nodes_after > 0, f"count={nodes_after}")

        # delete clears KB + row
        r = c.delete(f"/files/{txt_fid}", headers=H)
        check("txt delete 204", r.status_code == 204, str(r.status_code))
        r = c.get(f"/files/{txt_fid}", headers=H)
        check("deleted file -> 404", r.status_code == 404, str(r.status_code))
        nodes_post_delete = neo4j_label_count(cid)
        check("KB nodes dropped after delete", nodes_post_delete < nodes_after, f"{nodes_after}->{nodes_post_delete}")

        # ── 2) DOCX via live Docling ──
        doc = Document()
        doc.add_heading("Threads and Concurrency", level=1)
        doc.add_paragraph("A thread is the smallest unit of CPU scheduling. Threads within a process share memory.")
        doc.add_paragraph("A mutex protects a critical section so only one thread enters at a time.")
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        r = c.post(f"/files/upload/{fid_folder}", headers=H, files={"upload": ("threads.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        check("docx upload 201", r.status_code == 201, str(r.status_code))
        docx_fid = r.json()["id"]
        print("  ingesting docx (Docling — may take a minute)…")
        st = poll_status(c, H, docx_fid, timeout=300)
        check("docx -> completed (Docling)", st == "completed", st)

    print(f"\n{_passed} passed, {_failed} failed")
    return 1 if _failed else 0


if __name__ == "__main__":
    sys.exit(main())
